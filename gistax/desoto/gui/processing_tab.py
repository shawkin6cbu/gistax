import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
from docx import Document

class ProcessingTab(ttk.Frame):
    def __init__(self, parent, shared_data):
        super().__init__(parent, padding=10)
        self.shared_data = shared_data

        # Main layout frames
        left_frame = ttk.Frame(self)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 5))
        right_frame = ttk.Frame(self)
        right_frame.pack(side="right", fill="both", expand=True, padx=(5, 0))

        # --- Property Information ---
        prop_frame = ttk.LabelFrame(left_frame, text="Property Information", padding=10)
        prop_frame.pack(fill="x", pady=(0, 10))
        self.pin_var = tk.StringVar()
        self.address_var = tk.StringVar()
        self.owner_var = tk.StringVar()
        self.city_var = tk.StringVar()
        self.legal_desc_var = tk.StringVar()
        self._create_entry_row(prop_frame, "PIN:", self.pin_var, 0)
        self._create_entry_row(prop_frame, "Address:", self.address_var, 1)
        self._create_entry_row(prop_frame, "Owner:", self.owner_var, 2)
        self._create_entry_row(prop_frame, "City/State/ZIP:", self.city_var, 3)
        self._create_entry_row(prop_frame, "Legal Desc:", self.legal_desc_var, 4)

        # --- Tax Information ---
        tax_frame = ttk.LabelFrame(left_frame, text="Tax Information", padding=10)
        tax_frame.pack(fill="x", pady=(0, 10))
        self.tax_2024_total_var = tk.StringVar()
        self.tax_2024_paid_var = tk.StringVar()
        self.tax_2025_est_var = tk.StringVar()
        self._create_entry_row(tax_frame, "2024 Total:", self.tax_2024_total_var, 0)
        self._create_entry_row(tax_frame, "2024 Paid:", self.tax_2024_paid_var, 1)
        self._create_entry_row(tax_frame, "2025 Estimated:", self.tax_2025_est_var, 2)

        # --- Title Chain Summary ---
        title_frame = ttk.LabelFrame(left_frame, text="Title Chain Summary", padding=10)
        title_frame.pack(fill="x", pady=(0, 10))
        self.title_summary_var = tk.StringVar(value="Not loaded yet.")
        ttk.Label(title_frame, textvariable=self.title_summary_var, wraplength=250).pack(anchor="w")
        ttk.Button(title_frame, text="View Details", command=self.view_title_details).pack(pady=(5,0))

        # --- Document Details ---
        doc_details_frame = ttk.LabelFrame(right_frame, text="Document Details", padding=10)
        doc_details_frame.pack(fill="x", pady=(0, 10))
        self.lender_var = tk.StringVar()
        self.borrower_var = tk.StringVar()
        self.loan_amount_var = tk.StringVar()
        self.writer_var = tk.StringVar()
        self.date_var = tk.StringVar()
        self.notes_var = tk.StringVar()
        self._create_entry_row(doc_details_frame, "Lender:", self.lender_var, 0)
        self._create_entry_row(doc_details_frame, "Borrower:", self.borrower_var, 1)
        self._create_entry_row(doc_details_frame, "Loan Amount:", self.loan_amount_var, 2)
        self._create_entry_row(doc_details_frame, "Writer:", self.writer_var, 3)
        self._create_entry_row(doc_details_frame, "Date:", self.date_var, 4)
        self._create_entry_row(doc_details_frame, "Notes:", self.notes_var, 5)

        # --- Document Generation ---
        doc_gen_frame = ttk.LabelFrame(right_frame, text="Document Generation", padding=10)
        doc_gen_frame.pack(fill="x", pady=(0, 10))
        self.output_path_var = tk.StringVar()
        self._create_entry_row(doc_gen_frame, "Output Path:", self.output_path_var, 0, browse_btn=True)
        self.generate_btn = ttk.Button(doc_gen_frame, text="Generate Document", command=self.generate_document)
        self.generate_btn.grid(row=1, column=1, sticky="w", pady=(10,0))
        self.progress = ttk.Progressbar(doc_gen_frame, mode="indeterminate")
        self.progress.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(5,0))

        # --- Load from Tabs Button ---
        load_button = ttk.Button(self, text="Load From Tabs", command=self.load_from_tabs)
        load_button.pack(side="bottom", pady=10)

    def _create_entry_row(self, parent, label_text, var, row, browse_btn=False):
        ttk.Label(parent, text=label_text).grid(row=row, column=0, sticky="e", padx=(0, 5), pady=2)
        entry = ttk.Entry(parent, textvariable=var, width=30)
        entry.grid(row=row, column=1, sticky="ew", padx=(0, 5 if browse_btn else 0), pady=2)
        if browse_btn:
            ttk.Button(parent, text="...", width=3, command=self.browse_output).grid(row=row, column=2, pady=2)
        parent.columnconfigure(1, weight=1)

    def load_from_tabs(self):
        self.pin_var.set(self.shared_data.get_data("parcel_pin"))
        self.address_var.set(self.shared_data.get_data("parcel_address"))
        self.owner_var.set(self.shared_data.get_data("parcel_owner"))
        self.city_var.set(self.shared_data.get_data("parcel_city_state_zip"))
        self.legal_desc_var.set(self.shared_data.get_data("parcel_legal_description"))
        self.tax_2024_total_var.set(self.shared_data.get_data("tax_2024_total"))

        results = self.shared_data.get_data("title_chain_results")
        if results:
            self.title_summary_var.set(f"{len(results)} vesting deeds found.")
        else:
            self.title_summary_var.set("No title chain data found.")

    def view_title_details(self):
        results = self.shared_data.get_data("title_chain_results")
        if not results:
            messagebox.showinfo("Title Chain Details", "No title chain data has been loaded.")
            return

        details_win = tk.Toplevel(self)
        details_win.title("Title Chain Details")
        details_win.geometry("600x400")

        cols = ("Date", "Grantor", "Grantee", "Instrument", "Book-Page")
        tree = ttk.Treeview(details_win, columns=cols, show="headings")
        for col in cols:
            tree.heading(col, text=col)
            tree.column(col, width=110, anchor="w")
        tree.pack(fill="both", expand=True, padx=10, pady=10)

        for deed in results:
            tree.insert("", "end", values=(deed.date_string, deed.grantor, deed.grantee, deed.instrument, deed.book_page))

    def browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Save Document As",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
        )
        if path:
            self.output_path_var.set(path)

    def generate_document(self):
        output_path = self.output_path_var.get().strip()
        if not output_path:
            messagebox.showerror("Error", "Please specify an output path.")
            return

        self.generate_btn.config(state="disabled")
        self.progress.start()

        threading.Thread(target=self._generate_document_thread, args=(output_path,), daemon=True).start()

    def _generate_document_thread(self, output_path):
        try:
            success, msg = self._create_full_document(output_path)
            if success:
                self.after(0, lambda: messagebox.showinfo("Success", msg))
            else:
                self.after(0, lambda: messagebox.showerror("Error", msg))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", f"An unexpected error occurred: {e}"))
        finally:
            self.after(0, self.progress.stop)
            self.after(0, lambda: self.generate_btn.config(state="normal"))

    def get_template_path(self):
        try:
            # Assumes templates are in a 'templates' folder sibling to the 'desoto' folder
            current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            path = os.path.join(current_dir, "templates", "td_tmplt2.docx")
            return path if os.path.exists(path) else None
        except Exception:
            return None

    def _create_full_document(self, output_path):
        template_path = self.get_template_path()
        if not template_path:
            return False, "Template file 'td_tmplt2.docx' not found."

        doc = Document(template_path)

        # --- Gather all data ---
        placeholders = {
            "{PARCEL_NUMBER}": self.pin_var.get(),
            "{PROP_ADDR}": self.address_var.get(),
            "{OWNER}": self.owner_var.get(),
            "{CITY_STATE_ZIP}": self.city_var.get(),
            "{LEGAL_DESC}": self.legal_desc_var.get(),
            "{TAX_2024_TOTAL}": self.tax_2024_total_var.get(),
            "{TAX_2024_PAID}": self.tax_2024_paid_var.get(),
            "{TAX_2025_EST}": self.tax_2025_est_var.get(),
            "{LENDER}": self.lender_var.get(),
            "{BORROWER}": self.borrower_var.get(),
            "{LOAN_AMOUNT}": self.loan_amount_var.get(),
            "{WRITER}": self.writer_var.get(),
            "{DATE}": self.date_var.get(),
            "{NOTES}": self.notes_var.get(),
        }

        # --- Replace in paragraphs ---
        for p in doc.paragraphs:
            for key, value in placeholders.items():
                if key in p.text:
                    # Replace while preserving style
                    inline = p.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, value)
                            inline[i].text = text

        # --- Replace in tables ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in placeholders.items():
                            if key in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, value)
                                        inline[i].text = text

        # --- Fill Title Chain Table ---
        chain_deeds = self.shared_data.get_data("title_chain_results")
        chain_table = None
        for table in doc.tables:
            header_text = ' '.join([cell.text.upper() for cell in table.rows[0].cells])
            if 'GRANTOR' in header_text and 'GRANTEE' in header_text:
                chain_table = table
                break

        if chain_table:
            # Clear existing data rows (from row 1 downwards)
            while len(chain_table.rows) > 1:
                chain_table._element.remove(chain_table.rows[-1]._element)

            if chain_deeds:
                for deed in chain_deeds:
                    row_cells = chain_table.add_row().cells
                    row_cells[0].text = deed.grantor.upper()
                    row_cells[1].text = deed.grantee.upper()
                    row_cells[2].text = deed.instrument.upper()
                    row_cells[3].text = deed.date_string
                    row_cells[4].text = deed.book_page
            else:
                row_cells = chain_table.add_row().cells
                row_cells[0].text = "No vesting deeds found"

        doc.save(output_path)
        return True, f"Document successfully generated at:\n{output_path}"
