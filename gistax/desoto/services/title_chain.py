
import re
import PyPDF2
import fitz  # PyMuPDF - alternative PDF reader
from datetime import datetime, timedelta
from dataclasses import dataclass
from typing import List, Optional
from docx import Document
import os

@dataclass
class ChainEntry:
    date: datetime
    date_string: str
    grantor: str
    grantee: str
    instrument: str
    book_page: str
    remark: str = ""
    is_vesting: bool = False
    line: str = ""

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF using multiple methods for better reliability."""
    text = ""
    
    # Try PyMuPDF first (usually better text extraction)
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        if text.strip():
            return text
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}")
    
    # Fall back to PyPDF2
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
    except Exception as e:
        print(f"PyPDF2 extraction failed: {e}")
        
    return text

def parse_date(date_str: str) -> Optional[datetime]:
    """Parse MM/DD/YYYY format date string."""
    try:
        parts = date_str.split('/')
        if len(parts) == 3:
            month = int(parts[0])
            day = int(parts[1]) 
            year = int(parts[2])
            return datetime(year, month, day)
    except (ValueError, IndexError):
        pass
    return None

def is_vesting_deed(instrument: str) -> bool:
    """Determine if an instrument is a vesting deed."""
    vesting_deed_types = [
        'WARRANTY DEED',
        'SPECIAL WARRANTY DEED', 
        'QUITCLAIM DEED',
        'DEED',
        'GRANT DEED',
        'BARGAIN AND SALE DEED',
        'ASSUMPTION WARRANTY DEED',
        'CORRECTION DEED',
        'EXECUTOR\'S DEED',
        'ADMINISTRATOR\'S DEED',
        'TRUSTEE\'S DEED',
        'SHERIFF\'S DEED',
        'TAX DEED',
        'COMMISSIONER\'S DEED',
        'SPECIAL DEED',
        'BENEFICIARY DEED'
    ]
    
    non_vesting_types = [
        'DEED OF TRUST',
        'MORTGAGE',
        'ASSIGNMENT OF LEASES',
        'ASSIGNMENT OF RENTS', 
        'ASSIGNMENT OF INCOME',
        'ASSIGNMENT OF LEASES, RENTS AND INCOME',
        'ASSIGNMENT OF LEASES, REENTS AND INCOME',  # Handle typos
        'UCC FINANCING STATEMENT',
        'SATISFACTION',
        'RELEASE',
        'SUBORDINATION',
        'MODIFICATION',
        'EXTENSION',
        'LIS PENDENS',
        'NOTICE OF DEFAULT',
        'AFFIDAVIT',
        'EASEMENT',
        'RIGHT OF WAY'
    ]
    
    upper_instrument = instrument.upper().strip()
    
    # First check if it's explicitly a non-vesting type
    for non_vesting in non_vesting_types:
        if non_vesting in upper_instrument:
            return False
    
    # Then check if it's a vesting type
    for vesting in vesting_deed_types:
        if vesting in upper_instrument:
            return True
            
    return False

def preprocess_chain_text(text: str) -> str:
    """Preprocess text to handle multi-line entries."""
    lines = text.split('\n')
    processed_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Check if this line starts with a date (beginning of an entry)
        if re.match(r'^\d{2}/\d{2}/\d{4}', line):
            # This is a date line, combine with following non-date lines
            combined = line
            i += 1
            
            # Continue combining lines until we hit another date line, separator, or end
            while i < len(lines):
                next_line = lines[i].strip()
                
                # Stop if we hit another date line, separator, or specific patterns
                if (re.match(r'^\d{2}/\d{2}/\d{4}', next_line) or 
                    not next_line or 
                    '****' in next_line or
                    'FILED' in next_line or
                    'NAME CERTIFICATION' in next_line or
                    'SELLER' in next_line or
                    'BUYER' in next_line or
                    'Information to follow' in next_line):
                    break
                    
                combined += ' ' + next_line
                i += 1
            
            processed_lines.append(combined)
            continue
        
        processed_lines.append(line)
        i += 1
    
    return '\n'.join(processed_lines)

def parse_chain_text(text: str) -> List[ChainEntry]:
    """Parse chain of title text into structured entries using column positions."""
    lines = text.split('\n')
    entries = []
    
    # Find the header line with column positions
    header_idx = None
    col_positions = {}
    
    for i, line in enumerate(lines):
        # Look for the header line
        if 'GRANTOR' in line and 'GRANTEE' in line and 'INSTRUMENT' in line:
            header_idx = i
            # Get column positions from header
            col_positions = {
                'grantor': line.find('GRANTOR'),
                'grantee': line.find('GRANTEE'),
                'instrument': line.find('INSTRUMENT'),
                'dated': line.find('DATED'),
                'recording': line.find('RECORDING')
            }
            break
    
    if not header_idx or not col_positions:
        # Fallback to old regex method if no table found
        return parse_chain_text_regex_fallback(text)
    
    # Process table data
    in_table = False
    current_entry_lines = []
    
    for i in range(header_idx + 1, len(lines)):
        line = lines[i]
        
        # Start of table data (separator line)
        if '---' in line or '***' in line:
            if not in_table:
                in_table = True
                continue
            else:
                # End of table
                if current_entry_lines:
                    entry = parse_table_entry(current_entry_lines, col_positions)
                    if entry:
                        entries.append(entry)
                break
        
        if not in_table:
            continue
        
        # Empty line = end of current entry
        if not line.strip():
            if current_entry_lines:
                entry = parse_table_entry(current_entry_lines, col_positions)
                if entry:
                    entries.append(entry)
                current_entry_lines = []
            continue
        
        # Add line to current entry
        current_entry_lines.append(line)
    
    return entries

def parse_table_entry(lines: List[str], col_positions: dict) -> ChainEntry:
    """Parse a multi-line table entry using column positions."""
    if not lines:
        return None
    
    # Extract data from each column
    grantor_parts = []
    grantee_parts = []
    instrument_parts = []
    date_str = ""
    book_page = ""
    
    for line in lines:
        # Extract grantor
        if col_positions.get('grantor', -1) >= 0:
            start = col_positions['grantor']
            end = col_positions.get('grantee', len(line))
            text = line[start:end].strip()
            if text:
                grantor_parts.append(text)
        
        # Extract grantee
        if col_positions.get('grantee', -1) >= 0:
            start = col_positions['grantee']
            end = col_positions.get('instrument', len(line))
            text = line[start:end].strip()
            if text:
                grantee_parts.append(text)
        
        # Extract instrument
        if col_positions.get('instrument', -1) >= 0:
            start = col_positions['instrument']
            end = col_positions.get('dated', len(line))
            text = line[start:end].strip()
            if text:
                instrument_parts.append(text)
        
        # Extract date (usually only on first line)
        if not date_str and col_positions.get('dated', -1) >= 0:
            start = col_positions['dated']
            end = col_positions.get('recording', len(line))
            text = line[start:end].strip()
            if text and re.match(r'\d{2}/\d{2}/\d{4}', text):
                date_str = text
        
        # Extract recording/book-page (usually only on first line)
        if not book_page and col_positions.get('recording', -1) >= 0:
            start = col_positions['recording']
            text = line[start:].strip()
            if text and re.match(r'[\w\d]+-[\w\d]+', text):
                book_page = text
    
    # Combine multi-line fields
    grantor = ' '.join(grantor_parts).strip()
    grantee = ' '.join(grantee_parts).strip()
    instrument = ' '.join(instrument_parts).strip()
    
    # Create entry if we have minimum required fields
    if date_str and (grantor or grantee) and book_page:
        parsed_date = parse_date(date_str)
        if parsed_date:
            return ChainEntry(
                date=parsed_date,
                date_string=date_str,
                grantor=grantor,
                grantee=grantee,
                instrument=instrument,
                book_page=book_page,
                remark="",
                is_vesting=is_vesting_deed(instrument),
                line=' '.join(lines)
            )
    
    return None

def parse_chain_text_regex_fallback(text: str) -> List[ChainEntry]:
    """Original regex-based parsing as fallback for non-table formats."""
    # Your original preprocess and regex code here
    text = preprocess_chain_text(text)
    lines = text.split('\n')
    entries = []
    
    for line in lines:
        line = line.strip()
        
        # Skip headers, separators, and metadata
        skip_patterns = [
            r'FILED.*GRANTOR.*GRANTEE.*INSTRUMENT',
            r'^\*+',
            r'CHAIN OF TITLE',
            r'File No\.',
            r'NAME CERTIFICATION',
            r'SELLER\s+.*\s+BUYER',
            r'OWNER:',
            r'For further information',
            r'Certified to:',
            r'New Certification Date:',
            r'By:.*',
            r'INFORMATION TO FOLLOW',
            r'^\s*$'
        ]
        
        if any(re.search(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
            continue
        
        # Your original regex patterns
        patterns = [
            r'^(\d{2}/\d{2}/\d{4})\s+(.+?)\s+(.+?)\s+((?:[\w\s]+(?:DEED|TRUST|ASSIGNMENT|MORTGAGE|UCC|SATISFACTION|RELEASE|SUBORDINATION|MODIFICATION|EXTENSION|LIS PENDENS|NOTICE|AFFIDAVIT|EASEMENT)[\w\s]*)|(?:P\s+\d+-\d+))\s+([A-Z]?\s*\d+-\d+|\w+-\w+)(?:\s+(.*))?$',
            r'^(\d{2}/\d{2}/\d{4})\s+(.+?)\s+(WARRANTY DEED|DEED OF TRUST|QUITCLAIM DEED|SPECIAL WARRANTY DEED|DEED)\s+(\d+-\d+)(?:\s+(.*))?$'
        ]
        
        match = None
        for pattern in patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                break
        
        if match:
            groups = match.groups()
            
            if len(groups) >= 5:
                if len(groups) == 6:
                    date_str, grantor, grantee, instrument, book_page, remark = groups
                else:
                    date_str, combined_names, instrument, book_page, remark = groups
                    name_parts = combined_names.split()
                    if len(name_parts) >= 2:
                        mid_point = len(name_parts) // 2
                        grantor = ' '.join(name_parts[:mid_point])
                        grantee = ' '.join(name_parts[mid_point:])
                    else:
                        grantor = combined_names
                        grantee = "UNKNOWN"
                
                parsed_date = parse_date(date_str)
                
                if parsed_date:
                    entry = ChainEntry(
                        date=parsed_date,
                        date_string=date_str,
                        grantor=grantor.strip(),
                        grantee=grantee.strip(), 
                        instrument=instrument.strip(),
                        book_page=book_page.strip(),
                        remark=remark.strip() if remark else "",
                        is_vesting=is_vesting_deed(instrument.strip()),
                        line=line
                    )
                    entries.append(entry)
    
    return entries

def get_24_month_chain(entries: List[ChainEntry], processing_date: datetime = None) -> List[ChainEntry]:
    """
    Get chain of title covering AT LEAST the last 24 months.
    
    Logic:
    1. Find all vesting deeds
    2. Start from most recent and work backwards
    3. Include deeds until we cover at least 24 months from processing date
    4. If no deeds within 24 months, include the most recent vesting deed
    """
    if processing_date is None:
        processing_date = datetime.now()
        
    cutoff_date = processing_date - timedelta(days=730)  # 24 months ≈ 730 days
    
    # Get all vesting deeds, sorted by date (newest first)
    vesting_deeds = [entry for entry in entries if entry.is_vesting]
    vesting_deeds.sort(key=lambda x: x.date, reverse=True)
    
    if not vesting_deeds:
        return []
    
    result = []
    
    # If the most recent deed is within 24 months, include it and work backwards
    if vesting_deeds[0].date >= cutoff_date:
        # Start with the most recent deed
        result.append(vesting_deeds[0])
        
        # Work backwards to ensure we cover at least 24 months
        for deed in vesting_deeds[1:]:
            result.append(deed)
            
            # Check if we now cover at least 24 months
            earliest_date = min(deed.date for deed in result)
            if earliest_date <= cutoff_date:
                break
    else:
        # No deeds within 24 months, so include the most recent vesting deed
        result.append(vesting_deeds[0])
    
    # Sort result by date (newest first for display)
    result.sort(key=lambda x: x.date, reverse=True)
    return result

def create_title_document(chain_deeds: List[ChainEntry], output_path: str, template_path: str = None) -> bool:
    """Create title document with 24-month chain."""
    try:
        if template_path and os.path.exists(template_path):
            # Use provided template - this preserves the entire document
            doc = Document(template_path)
        else:
            # Create basic document
            doc = Document()
            doc.add_heading('TWENTY-FOUR MONTH CHAIN OF TITLE', level=1)
        
        # Find the chain table - look for any table that has chain-related headers
        chain_table = None
        
        # Search all tables for one with chain-related headers
        for table in doc.tables:
            if len(table.rows) > 0:
                header_row = table.rows[0]
                header_text = ' '.join([cell.text.upper() for cell in header_row.cells])
                
                # Look for typical chain table headers
                if any(keyword in header_text for keyword in ['GRANTOR', 'GRANTEE', 'INSTRUMENT', 'DATED', 'RECORDING']):
                    chain_table = table
                    break
        
        # If no suitable table found, look for tables near "TWENTY-FOUR MONTH CHAIN" text
        if not chain_table:
            for paragraph in doc.paragraphs:
                if "TWENTY-FOUR MONTH CHAIN" in paragraph.text.upper() or "24 MONTH CHAIN" in paragraph.text.upper():
                    # Look for tables after this paragraph
                    para_element = paragraph._element
                    parent = para_element.getparent()
                    
                    # Find this paragraph's position and look for subsequent tables
                    found_para = False
                    for element in parent:
                        if found_para and element.tag.endswith('tbl'):
                            # Found a table after the chain paragraph
                            for table in doc.tables:
                                if table._element == element:
                                    chain_table = table
                                    break
                            break
                        if element == para_element:
                            found_para = True
                    
                    if chain_table:
                        break
        
        # Create table if still not found (fallback)
        if not chain_table:
            chain_table = doc.add_table(rows=1, cols=5)
            chain_table.style = 'Table Grid'
            
            # Set headers
            header_cells = chain_table.rows[0].cells
            header_cells[0].text = 'GRANTOR:'
            header_cells[1].text = 'GRANTEE:'
            header_cells[2].text = 'INSTRUMENT'
            header_cells[3].text = 'DATED:'
            header_cells[4].text = 'RECORDING:'
        
        # Clear existing data rows (preserve header if it exists)
        rows_to_remove = []
        for i, row in enumerate(chain_table.rows):
            if i == 0:  # Keep first row as header
                continue
            rows_to_remove.append(row)
        
        # Remove data rows
        for row in rows_to_remove:
            chain_table._element.remove(row._element)
        
        # Add chain data
        if chain_deeds:
            for deed in chain_deeds:
                row = chain_table.add_row()
                cells = row.cells
                if len(cells) >= 5:
                    cells[0].text = deed.grantor.upper()
                    cells[1].text = deed.grantee.upper()
                    cells[2].text = deed.instrument.upper()
                    cells[3].text = deed.date_string
                    cells[4].text = deed.book_page
                else:
                    # Handle tables with different column counts
                    for i, text in enumerate([deed.grantor.upper(), deed.grantee.upper(), deed.instrument.upper(), deed.date_string, deed.book_page]):
                        if i < len(cells):
                            cells[i].text = text
        else:
            # Add empty row if no deeds found
            row = chain_table.add_row()
            if len(row.cells) > 0:
                row.cells[0].text = "No vesting deeds found"
        
        # Save the complete document (preserving all template content)
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating document: {e}")
        return False

def process_title_pdf(pdf_path: str, output_path: str, template_path: str = None) -> tuple[bool, str, List[ChainEntry]]:
    """
    Main function to process title PDF and create document.
    Returns: (success, message, chain_deeds)
    """
    try:
        # Extract text
        text = extract_text_from_pdf(pdf_path)
        if not text.strip():
            return False, "Could not extract text from PDF", []
        
        # Parse entries
        entries = parse_chain_text(text)
        if not entries:
            return False, "No chain of title entries found", []
        
        # Get 24-month chain
        chain_deeds = get_24_month_chain(entries)
        
        # Create document
        success = create_title_document(chain_deeds, output_path, template_path)
        
        if success:
            message = f"Document created successfully!\nFound {len(chain_deeds)} vesting deeds in 24-month chain."
            return True, message, chain_deeds
        else:
            return False, "Failed to create document", chain_deeds
            
    except Exception as e:
        return False, f"Error processing PDF: {str(e)}", []